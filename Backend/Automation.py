# File: modules/Automation.py (Đã nâng cấp hàm Content)

# --- 1. Import các thư viện cần thiết ---
from AppOpener import close, open as appopen
from webbrowser import open as webopen
from pywhatkit import search, playonyt
from dotenv import dotenv_values
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import google.generativeai as genai # <<< THÊM: Import Gemini
import webbrowser
import subprocess
import requests
import keyboard
import asyncio
import os
import PIL.Image
# --- 2. Load biến môi trường và Khởi tạo API ---
env_path = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '.env'))
env_vars = dotenv_values(env_path)

project_root = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
DATA_DIR = os.path.join(project_root, "Data") # Đây là đường dẫn tới Data/ ở gốc

# Lấy API Key của Gemini (giống file brain_engine.py)
GoogleAPIKey = env_vars.get("GEMINI_API_KEY")
Username = env_vars.get("User_name")

# Khởi tạo Gemini
if not GoogleAPIKey:
    print("❌ Lỗi: Không tìm thấy 'GEMINI_API_KEY' trong file .env (cần cho hàm Content)")
else:
    try:
        genai.configure(api_key=GoogleAPIKey)
        print("✅ Đã kết nối tới Gemini (Automation).")
    except Exception as e:
        print(f"❌ Lỗi khởi tạo Gemini Client (Automation): {e}")

# --- 3. Cấu hình (từ ảnh) ---
# (useragent, classes, professional_responses... giữ nguyên)
classes = [
    "Z0LcW", "gsrt vk_bk FZvtNb YmP-hnf", "pclqee", "tw-data-text tw-text-small tw-ta",
    "IZ6rdc", "05UR6d LTKOO", "vLqJ-w", "webanswers-webanswers_table__webanswers-table", "dDoNo ikb48b gsrt", "sXLAoe",
    "LWkfKe", "VQF4g", "qv3Wpe", "kno-rdesc", "SPZ26b"
] 
useragent = "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/100.0.4896.75 Safari/537.36"
professional_responses = [
    "Sự hài lòng của bạn là ưu tiên hàng đầu của tôi...",
    "Tôi luôn sẵn sàng phục vụ bạn...",
]
WEB_FALLBACK = {
    "zalo": "https://chat.zalo.me/",
    "facebook": "https://facebook.com",
    "chrome": "https://google.com",
    "spotify": "https://open.spotify.com",
    "telegram": "https://web.telegram.org",
    "messenger": "https://messenger.com",
    "youtube": "https://youtube.com",
    "canva": "https://www.canva.com/",
    "discord": "https://discord.com/app",
    "Gmail": "https://mail.google.com/",
}
def GoogleSearch(Topic):
    try:
        search(Topic) 
        return True 
    except Exception as e:
        print(f"❌ Lỗi GoogleSearch (pywhatkit): {e}")
        return False

def YouTubeSearch(Topic):
    try:
        Url4Search = f"https://www.youtube.com/results?search_query={Topic}"
        webopen(Url4Search) 
        return True
    except Exception as e:
        print(f"❌ Lỗi YouTubeSearch: {e}")
        return False

def PlayYouTube(query):
    try:
        playonyt(query) 
        return True
    except Exception as e:
        print(f"❌ Lỗi PlayYouTube (pywhatkit): {e}")
        return False

# --- B. Công cụ viết nội dung (ĐÃ NÂNG CẤP DÙNG GEMINI) ---
def Content(Topic):
    """Dùng AI (Gemini) để viết nội dung, lưu ra file .docx và mở bằng Word."""

    # --- A. Mở file Word ---
    def OpenWord(File):
        try:
            os.startfile(File)
        except Exception as e:
            print(f"❌ Lỗi mở Word: {e}")

    # --- B. Gọi AI Gemini để sinh nội dung ---
    def ContentWriterAI(prompt):
        try:
            model = genai.GenerativeModel('gemini-2.5-flash')
            full_prompt = (
                f"Bạn là một người viết nội dung chuyên nghiệp. "
                f"Người dùng '{Username}' yêu cầu bạn viết về chủ đề sau: \"{prompt}\".\n\n"
                "Hãy viết nội dung chi tiết, có cấu trúc rõ ràng (dùng các mục I., II., III., 1., 2.), "
                "bằng tiếng Việt. Không thêm lời chào hoặc giới thiệu."
            )
            print("🤖 (ContentWriter-Gemini) Đang viết nội dung Word...")
            response = model.generate_content(full_prompt)
            Answer = response.text.strip()
            return Answer

        except Exception as e:
            print(f"❌ Lỗi ContentWriterAI (Gemini): {e}")
            return f"Lỗi khi tạo nội dung: {e}"

    # --- C. Rearrange Paragraphs thông minh ---
    def rearrange_paragraphs(doc, text):
        lines = [l.strip() for l in text.split("\n") if l.strip()]
        paragraph_buffer = ""

        # Nhận diện các mẫu mục lớn / nhỏ
        new_section_pattern = re.compile(r"^(?:[IVXLCDM]+\.)|(?:\d+\.)|(?:[A-Z]\))|(?:\*\*|##)")

        for line in lines:
            if new_section_pattern.match(line):
                # Nếu có đoạn cũ => ghi vào doc
                if paragraph_buffer:
                    p = doc.add_paragraph(paragraph_buffer.strip())
                    p.style = doc.styles["Normal"]
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.style.font.name = "Times New Roman"
                    p.style.font.size = Pt(12)
                    paragraph_buffer = ""

                # In đậm nếu là mục lớn (I., II., 1.)
                if re.match(r"^(?:[IVXLCDM]+\.|\d+\.)", line):
                    p = doc.add_paragraph(line)
                    run = p.runs[0]
                    run.bold = True
                    p.style = doc.styles["Normal"]
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.style.font.name = "Times New Roman"
                    p.style.font.size = Pt(12)
                else:
                    paragraph_buffer += line + "\n"
            else:
                paragraph_buffer += " " + line

        # Ghi nốt đoạn cuối cùng
        if paragraph_buffer:
            p = doc.add_paragraph(paragraph_buffer.strip())
            p.style = doc.styles["Normal"]
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.style.font.name = "Times New Roman"
            p.style.font.size = Pt(12)

# --- D. Thực thi ---
    try:
        # Giữ Topic gốc (có dấu "?") để dùng cho tiêu đề và gọi AI
        Topic = str(Topic).replace("content ", "", 1).strip()
        ContentByAI = ContentWriterAI(Topic)

        # --- PHẦN SỬA LỖI ---
        # 1. Tạo một bản sao của Topic để làm sạch cho tên file
        #    Loại bỏ TẤT CẢ các ký tự không hợp lệ của Windows
        sanitized_topic = re.sub(r'[\\/*?:"<>|]', '', Topic)
        
        # 2. Chuyển đổi tên đã làm sạch thành tên file (không còn dấu "?")
        safe_filename = sanitized_topic.lower().replace(' ', '_') + ".docx"
        # Kết quả mong muốn: 'hình_lăng_trụ_đừng_là_gì.docx'
        # --- KẾT THÚC SỬA LỖI ---

        file_path = os.path.abspath(os.path.join(DATA_DIR, safe_filename))
        os.makedirs(DATA_DIR, exist_ok=True)

        doc = Document()
        # Dùng Topic gốc (có dấu "?") làm tiêu đề trong văn bản Word
        doc.add_heading(Topic, level=1) 

        rearrange_paragraphs(doc, ContentByAI)

        doc.save(file_path)
        print(f"✅ Đã lưu nội dung vào: {file_path}")

        OpenWord(file_path)
        return True

    except Exception as e:
        print(f"❌ Lỗi trong hàm Content: {e}")
        return False

# --- C. Công cụ Điều khiển Ứng dụng/Web (Giữ nguyên) ---
def OpenApp(app, sess_requests=None, html_content=None):
    """
    Hàm này mở App hoặc Web.
    Ưu tiên mở App Desktop, nếu thất bại, sẽ thử mở Web.
    """
    
    # Hàm lồng: Trích xuất link từ HTML (cho tìm kiếm web)
    def extract_links(html):
        if html is None: return []
        soup = BeautifulSoup(html, 'html.parser')
        links = soup.find_all('a', {'jsname': 'UWckNb'}) 
        return [link.get('href') for link in links]

    # 1. Ưu tiên mở link web nếu có HTML (từ Google Search)
    if html_content:
        try:
            links = extract_links(html_content)
            if links:
                webopen(links[0]) # Mở link đầu tiên tìm thấy
                return True
        except Exception as e:
            print(f"❌ Lỗi extract_links: {e}")
            # (Tiếp tục thử mở app bên dưới)
    
    # 2. Thử mở App Desktop
    app_key = app.lower().strip()
    try:
        print(f"Đang thử mở app: {app_key}")
        appopen(app_key, match_closest=True, output=True, throw_error=True)
        return True
    except Exception as e:
        # 3. LỖI (Ví dụ: ZALO is not running) -> Thử mở Web
        print(f"⚠️ Lỗi OpenApp: {e}")
        print(f"💡 Không tìm thấy app, thử mở web fallback...")
        
        if app_key in WEB_FALLBACK:
            try:
                web_url = WEB_FALLBACK[app_key]
                print(f"Đang mở web: {web_url}")
                webopen(web_url) # Dùng webbrowser để mở
                return True
            except Exception as e_web:
                print(f"❌ Lỗi mở web fallback: {e_web}")
                return False
        else:
            print(f"🤷 Không có web fallback cho: {app_key}")
            return False

def CloseApp(app):
    try:
        # Map tên thân thiện sang tên file .exe (chuẩn hóa tất cả là .exe)
        aliases = {
            "microsoft word": "WINWORD.EXE",
            "word": "WINWORD.EXE",
            "microsoft excel": "EXCEL.EXE",
            "excel": "EXCEL.EXE",
            "powerpoint": "POWERPNT.EXE",
            "chrome": "chrome.exe",
            "notepad": "notepad.exe",
            "visual studio 2022": "devenv.exe",
            "wps office": "wps.exe",
            "dev-c++": "devcpp.exe",
            "dev c++": "devcpp.exe",
            "devcpp": "devcpp.exe",
        }

        app_key = (app or "").lower().strip()
        exe_name = aliases.get(app_key, app_key)

        # ensure .exe suffix
        if not exe_name.lower().endswith('.exe'):
            exe_name = exe_name + '.exe'

        print(f"Đang cố gắng đóng: {exe_name}")

        # Thử kill bằng taskkill theo tên exe
        try:
            res = subprocess.run(["taskkill", "/f", "/im", exe_name], capture_output=True, text=True)
            out = (res.stdout or "").strip()
            err = (res.stderr or "").strip()
            if res.returncode == 0:
                print(f"✅ Đã đóng {app} ({exe_name})")
                return True
            else:
                # rõ ràng thất bại, log lỗi để debug
                print(f"⚠️ taskkill trả về mã {res.returncode}. stdout: {out}; stderr: {err}")
                return False
        except Exception as e:
            print(f"❌ Lỗi khi chạy taskkill: {e}")
            return False

    except Exception as e:
        print(f"❌ Lỗi CloseApp: {e}")
        return False

# --- D. Công cụ Điều khiển Hệ thống (Giữ nguyên) ---
def System(command):
    # (Code hàm này giữ nguyên)
    def mute():
        keyboard.press_and_release("volume mute")
    def unmute():
        keyboard.press_and_release("volume mute")
    def volume_up():
        keyboard.press_and_release("volume up")
    def volume_down():
        keyboard.press_and_release("volume down")
    try:
        cmd = command.lower().strip()
        if cmd == "mute" or "tắt tiếng" in cmd:
            mute()
        elif cmd == "unmute" or "bật tiếng" in cmd:
            unmute()
        elif "volume up" in cmd or "tăng âm lượng" in cmd:
            volume_up()
        elif "volume down" in cmd or "giảm âm lượng" in cmd:
            volume_down()
        else:
            print(f"⚠️ Không hiểu lệnh hệ thống: {cmd}")
            return False
        return True
    except Exception as e:
        print(f"❌ Lỗi System: {e}")
        return False

# --- E. Công cụ Scrape Web (Giữ nguyên) ---
def search_google(query, sess):
    # (Code hàm này giữ nguyên)
    url = f"https://www.google.com/search?q={query}"
    headers = {"User-Agent": useragent}
    response = sess.get(url, headers=headers)
    if response.status_code == 200:
        return response.text 
    else:
        print("Không thể lấy kết quả tìm kiếm.")
        return None

# --- 5. BỘ ĐIỀU PHỐI (Giữ nguyên) ---
async def TranslateAndExecute(commands: list[str]):
    # (Code hàm này giữ nguyên, vì nó đã được Việt hóa ở lần trước)
    funcs = [] 
    for command in commands:
        command = command.strip() 
        if command.startswith("mở "):
            app_name = command.removeprefix("mở ") 
            fun = asyncio.to_thread(OpenApp, app_name)
            funcs.append(fun)
        elif command.startswith("đóng "):
            app_name = command.removeprefix("đóng ")
            fun = asyncio.to_thread(CloseApp, app_name)
            funcs.append(fun)
        elif command.startswith("phát "):
            song_name = command.removeprefix("phát ")
            fun = asyncio.to_thread(PlayYouTube, song_name)
            funcs.append(fun)
        elif command.startswith("nội dung "):
            topic = command.removeprefix("nội dung ")
            fun = asyncio.to_thread(Content, topic)
            funcs.append(fun)
        elif command.startswith("tìm google "):
            query = command.removeprefix("tìm google ")
            fun = asyncio.to_thread(GoogleSearch, query)
            funcs.append(fun)
        elif command.startswith("tìm youtube "):
            query = command.removeprefix("tìm youtube ")
            fun = asyncio.to_thread(YouTubeSearch, query)
            funcs.append(fun)
        elif command.startswith("hệ thống "):
            sys_cmd = command.removeprefix("hệ thống ")
            fun = asyncio.to_thread(System, sys_cmd)
            funcs.append(fun)
        elif command.startswith("chung "):
            pass 
        elif command.startswith("thời gian thực "):
            pass 
        elif command.startswith("thoát"):
            pass
        else:
            print(f"⚠️ (Automation) Không tìm thấy hàm nào cho lệnh: {command}")
    results = await asyncio.gather(*funcs)
    for result in results:
        if isinstance(result, str):
            yield result
        else:
            yield result

async def Automation(commands: list[str]):
    # (Code hàm này giữ nguyên)
    async for result in TranslateAndExecute(commands):
        pass 
    return True